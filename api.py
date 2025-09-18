import json, os, sys, subprocess
from datetime import datetime
from openpyxl import Workbook, load_workbook
from openpyxl.utils import range_boundaries, get_column_letter
from docx import Document

SCHEMA_PATH = os.path.join('forms', 'schema.json')
TEMPLATES_DIR = 'templates'


class LocalAPI:
    def get_schema(self):
        with open(SCHEMA_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)

    def save_schema(self, schema):
        if not isinstance(schema, dict) or 'fields' not in schema:
            raise ValueError('Invalid schema: missing fields')
        os.makedirs(os.path.dirname(SCHEMA_PATH), exist_ok=True)
        with open(SCHEMA_PATH, 'w', encoding='utf-8') as f:
            json.dump(schema, f, ensure_ascii=False, indent=2)
        return True

    def generate_xlsx(self, data, folder):
        """Generate or append to XLSX, honoring an optional template.

        Behavior:
        - If an output file exists in the target folder, append to it.
        - Else if a known template exists under templates/, copy/initialize from it and write.
        - Else fall back to the simple header+append behavior.

        Mapping:
        - If a header row is present (plain sheet or table), map values by matching
          schema field label (preferred) or key to header text.
        - If no header row exists, create one from provided data keys (legacy behavior).
        """
        os.makedirs(folder or '.', exist_ok=True)
        xlsx_path = os.path.join(folder or '.', 'output.xlsx')

        # Discover an XLSX template deterministically from the templates folder.
        # Policy:
        # - If exactly one .xlsx exists, use it.
        # - If multiple exist, prefer common names; else choose lexicographically first for stability.
        template_path = None
        try:
            xlsx_files = [f for f in os.listdir(TEMPLATES_DIR) if f.lower().endswith('.xlsx')]
        except FileNotFoundError:
            xlsx_files = []
        if len(xlsx_files) == 1:
            template_path = os.path.join(TEMPLATES_DIR, xlsx_files[0])
        elif len(xlsx_files) > 1:
            preferred = [
                'xlsx_template.xlsx',
                'table.xlsx',
                'single-row-template.xlsx',
                'sddl-template-header-fields.xlsx',
            ]
            lower_map = {f.lower(): f for f in xlsx_files}
            picked = next((lower_map[name] for name in preferred if name in lower_map), None)
            if picked is None:
                picked = sorted(xlsx_files, key=lambda s: s.lower())[0]
            template_path = os.path.join(TEMPLATES_DIR, picked)

        # Load schema for labels/keys mapping (best-effort; fallback to data keys)
        try:
            schema = self.get_schema()
        except Exception:
            schema = {}

        fields = []
        if isinstance(schema, dict):
            # support either flat fields or sections[*].fields
            if 'sections' in schema and isinstance(schema['sections'], list):
                for section in schema['sections']:
                    fields.extend(section.get('fields', []))
            if not fields:
                fields = schema.get('fields', []) or []

        def load_workbook_from_template_or_existing():
            if os.path.exists(xlsx_path):
                return load_workbook(xlsx_path)
            if template_path and os.path.exists(template_path):
                # Start from template
                wb_t = load_workbook(template_path)
                # Save immediately to establish the output file with template formatting
                wb_t.save(xlsx_path)
                return load_workbook(xlsx_path)
            # Fallback new workbook
            return Workbook()

        wb = load_workbook_from_template_or_existing()
        ws = wb.active

        # Try to locate a table to append into. If found, capture its bounds and headers.
        table_obj = None
        table_min_col = table_min_row = table_max_col = table_max_row = None
        headers = []
        if hasattr(ws, 'tables') and ws.tables:
            # Prefer a table named FormData; otherwise pick the first
            preferred = ws.tables.get('FormData') if isinstance(ws.tables, dict) else None
            if preferred is None:
                # openpyxl 3.1 uses dict mapping; but support generic iteration
                try:
                    preferred = next(iter(ws.tables.values()))
                except Exception:
                    preferred = None
            table_obj = preferred
            if table_obj is not None and getattr(table_obj, 'ref', None):
                min_col, min_row, max_col, max_row = range_boundaries(table_obj.ref)
                table_min_col, table_min_row, table_max_col, table_max_row = (
                    min_col, min_row, max_col, max_row
                )
                # Header row is the first row in the table range
                for c in range(min_col, max_col + 1):
                    headers.append(ws.cell(row=min_row, column=c).value)

        # If no table headers, fall back to first row as header
        if not headers and ws.max_row >= 1:
            headers = [cell.value for cell in ws[1]]
            table_min_col, table_min_row = 1, 1
            table_max_col = len(headers) if headers else 0

        # Create header row if still missing
        if not headers:
            # Legacy fallback: create headers from provided data order
            headers = list(data.keys())
            ws.append(headers)
            table_min_col, table_min_row = 1, 1
            table_max_col = len(headers)

        # Build a helper to get value by header name using label/key matching
        # Prefer field label match; then key; finally raw data key.
        label_by_key = {f.get('key'): f.get('label') for f in fields if isinstance(f, dict)}
        key_by_label = {v: k for k, v in label_by_key.items() if v}

        def value_for_header(h):
            if h is None:
                return ''
            # Normalize header to string for matching
            hs = str(h)
            # Prefer label match
            if hs in key_by_label:
                return data.get(key_by_label[hs], '')
            # Fallback: header equals key
            if hs in data:
                return data.get(hs, '')
            # Final fallback: try to match by label directly in values dict
            # (UI always sends keys; this is just extra safety.)
            for k, lbl in label_by_key.items():
                if lbl == hs and k in data:
                    return data.get(k, '')
            return ''

        # Determine target row index for append
        if table_obj is not None and table_min_col is not None:
            # Find next empty row in the first column of the table
            r = (table_min_row or 1) + 1  # first data row below header
            while ws.cell(row=r, column=table_min_col).value not in (None, ''):
                r += 1
            target_row = r
            start_col = table_min_col
        else:
            # Append on the worksheet after the last row, starting column 1
            target_row = ws.max_row + 1 if ws.max_row else 2
            start_col = 1

        # Write values aligned with headers
        if headers:
            for idx, header in enumerate(headers):
                ws.cell(row=target_row, column=start_col + idx, value=value_for_header(header))
        else:
            # Should not happen, but keep compatibility
            ws.append([data.get(k, '') for k in ws[1]])

        # Copy style from the previous data row if available for consistent formatting
        prev_row = target_row - 1
        if prev_row >= (table_min_row or 1) + 1 and (table_max_col or 0) >= (table_min_col or 1):
            for c in range(start_col, start_col + len(headers)):
                src = ws.cell(row=prev_row, column=c)
                dst = ws.cell(row=target_row, column=c)
                dst.font = src.font
                dst.fill = src.fill
                dst.alignment = src.alignment
                dst.number_format = src.number_format
                dst.border = src.border

        # If we appended into a table, extend its ref to include the new row
        if table_obj is not None and table_min_col is not None and headers:
            end_col = start_col + len(headers) - 1
            start_col_letter = get_column_letter(start_col)
            end_col_letter = get_column_letter(end_col)
            # Table ref includes header row at top
            new_ref = f"{start_col_letter}{table_min_row}:{end_col_letter}{target_row}"
            table_obj.ref = new_ref

        wb.save(xlsx_path)
        return os.path.abspath(xlsx_path)

    def generate_docx(self, data, folder):
        os.makedirs(folder or '.', exist_ok=True)
        docx_path = os.path.join(folder or '.', 'output.docx')
        template = os.path.join(TEMPLATES_DIR, 'docx_base_template.docx')
        if os.path.exists(template):
            doc = Document(template)
            for p in doc.paragraphs:
                for k, v in data.items():
                    p.text = p.text.replace(f'{{{{{k}}}}}', str(v))
        else:
            doc = Document()
            doc.add_heading('Form Output', level=1)
            for k, v in data.items():
                doc.add_paragraph(f"{k}: {v}")
        doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
        doc.save(docx_path)
        return os.path.abspath(docx_path)

    def open_folder(self, path):
        path = os.path.abspath(path or '.')
        if sys.platform.startswith('win'):
            os.startfile(path)  # type: ignore[attr-defined]
        elif sys.platform == 'darwin':
            subprocess.call(['open', path])
        else:
            subprocess.call(['xdg-open', path])
        return True

    # --- Schema utilities ---
    def _detect_xlsx_template(self):
        try:
            xlsx_files = [f for f in os.listdir(TEMPLATES_DIR) if f.lower().endswith('.xlsx')]
        except FileNotFoundError:
            xlsx_files = []
        if len(xlsx_files) == 1:
            return os.path.join(TEMPLATES_DIR, xlsx_files[0])
        if len(xlsx_files) > 1:
            preferred = [
                'xlsx_template.xlsx', 'table.xlsx', 'single-row-template.xlsx', 'sddl-template-header-fields.xlsx'
            ]
            lower_map = {f.lower(): f for f in xlsx_files}
            for name in preferred:
                if name in lower_map:
                    return os.path.join(TEMPLATES_DIR, lower_map[name])
            return os.path.join(TEMPLATES_DIR, sorted(xlsx_files, key=lambda s: s.lower())[0])
        return None

    def _headers_from_template(self):
        template = self._detect_xlsx_template()
        if not template or not os.path.exists(template):
            return []
        try:
            from openpyxl import load_workbook
            wb = load_workbook(template)
            ws = wb.active
            headers = [c.value if c.value is not None else '' for c in ws[1]]
            # trim trailing empties
            while headers and (headers[-1] is None or str(headers[-1]).strip() == ''):
                headers.pop()
            headers = [str(h).strip() for h in headers if str(h).strip()]
            return headers
        except Exception:
            return []

    def propose_schema_from_template(self):
        """Return a proposed schema based on the XLSX template headers.

        - Preserves existing field configs when labels/keys match headers.
        - Adds new text fields for new headers.
        - Keeps unmatched existing fields out of the active list but reports them in meta.unmapped.
        - Increments version.
        """
        headers = self._headers_from_template()
        current = {}
        try:
            current = self.get_schema()
        except Exception:
            current = {}

        existing_fields = current.get('fields', []) if isinstance(current, dict) else []
        by_label = {str(f.get('label', '')).strip().lower(): f for f in existing_fields if isinstance(f, dict)}
        by_key = {str(f.get('key', '')).strip().lower(): f for f in existing_fields if isinstance(f, dict)}

        def to_key(name: str) -> str:
            import re
            s = re.sub(r'[^0-9A-Za-z]+', '_', name).strip('_').lower()
            s = re.sub(r'_+', '_', s)
            return s or 'field'

        proposed_fields = []
        used = set()
        for h in headers:
            h_norm = str(h).strip()
            f = by_label.get(h_norm.lower()) or by_key.get(to_key(h_norm))
            if f:
                # clone to avoid mutating original
                nf = dict(f)
                nf['label'] = h_norm
                # keep existing key/type/options
                proposed_fields.append(nf)
                used.add(id(f))
            else:
                proposed_fields.append({
                    'key': to_key(h_norm),
                    'label': h_norm,
                    'type': 'text',
                    'required': False
                })

        # Unmapped existing fields (not in headers)
        unmapped = []
        for f in existing_fields:
            if id(f) not in used:
                unmapped.append({'key': f.get('key'), 'label': f.get('label')})

        # Version bump
        ver = current.get('version', 1)
        try:
            ver = int(ver) + 1
        except Exception:
            ver = 1

        proposed = {
            'title': current.get('title', 'FormGen'),
            'version': ver,
            'fields': proposed_fields,
            'xlsx': current.get('xlsx', {'row_strategy': 'single_row', 'email_delimiter': ';'}),
            'meta': { 'unmapped': unmapped, 'source': 'xlsx_template_headers' }
        }
        return proposed
